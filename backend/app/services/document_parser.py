"""文档解析服务：按页解析 PDF / DOCX / TXT，提取表格与目录，扫描页 OCR 降级。

核心设计：
- 逐页建立 DocumentPage 记录（禁止整篇正文存单字段）
- 按内容块建立 DocumentChunk（含标题层级、页码、token 数）
- 每页保留真实页码；DOCX 无真实分页时明确标记"段落位置"
- 文字过少或主要由图片组成的页标记为"疑似扫描页"→ 调用 OCR
- OCR 不可用只标记该页状态，不影响整体解析
"""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.adapters.factory import get_ocr_adapter
from app.core.logging import get_logger

logger = get_logger(__name__)

# 文档分类枚举（与 SourceDocument.doc_type 对齐）
DOC_TYPES = {
    "tender": "招标文件",
    "scoring": "评分办法",
    "construction": "施工组织设计",
    "profile": "项目概况",
    "schedule": "总进度计划",
    "special": "专项施工方案",
    "qualification": "企业资信及案例",
    "other": "其他资料",
}


@dataclass
class ParsedPage:
    page_number: int
    location_label: str
    raw_text: str
    cleaned_text: str
    markdown_text: str
    page_type: str = "text"  # text | table | scan | mixed
    extraction_method: str = "native"  # native | ocr | mixed
    ocr_status: str = "none"
    confidence: float | None = None
    tables: list[dict] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedChunk:
    page_start: int | None
    page_end: int | None
    heading_path: str | None
    content: str
    chunk_type: str = "paragraph"  # heading | paragraph | table
    metadata: dict = field(default_factory=dict)
    sequence: int = 0


@dataclass
class ParsedDocument:
    pages: list[ParsedPage]
    chunks: list[ParsedChunk]
    toc: list[dict]
    total_pages: int
    ocr_pages: int
    failed_pages: int
    table_count: int


def _clean_text(text: str) -> str:
    """清洗文本：合并空白、去除多余空行。"""
    text = text.replace("　", " ").replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _count_tokens(text: str) -> int:
    """粗略 token 数（中文按 1.5 字/token 估算 + 英文词）。"""
    if not text:
        return 0
    cjk = len(re.findall(r"[一-鿿]", text))
    words = len(re.findall(r"[A-Za-z0-9]+", text))
    return int(cjk / 1.5 + words)


def _suggest_doc_type_by_name(file_name: str) -> str:
    """根据文件名猜测文档分类。"""
    name = file_name.lower()
    if "评分" in file_name or "评标" in file_name or "办法" in file_name:
        return "scoring"
    if "施工组织设计" in file_name:
        return "construction"
    if "概况" in file_name or "概述" in file_name:
        return "profile"
    if "进度" in file_name or "计划" in file_name:
        return "schedule"
    if "方案" in file_name:
        return "special"
    if "资质" in file_name or "业绩" in file_name or "资信" in file_name or "案例" in file_name:
        return "qualification"
    return "tender"


# ============================================================
# PDF 解析
# ============================================================

def _parse_pdf(data: bytes) -> ParsedDocument:
    return _parse_pdf_from_source(io.BytesIO(data))


def _parse_pdf_from_source(source) -> ParsedDocument:
    """从文件路径或二进制流解析 PDF，避免把整个文件读进内存。"""
    import pdfplumber

    pages: list[ParsedPage] = []
    chunks: list[ParsedChunk] = []
    toc: list[dict] = []
    ocr_pages = 0
    failed_pages = 0
    table_count = 0
    source_bytes: bytes | None = None

    with pdfplumber.open(source) as pdf:
        total = len(pdf.pages)
        heading_stack: list[str] = []

        for idx, page in enumerate(pdf.pages, start=1):
            raw_text = page.extract_text() or ""
            tables = []
            try:
                for t in page.extract_tables() or []:
                    if t:
                        tables.append(_table_to_json(t))
            except Exception:
                pass

            has_tables = bool(tables)
            cleaned = _clean_text(raw_text)
            text_len = len(cleaned.strip())

            # 疑似扫描页：文字过少（<30字）或图片为主
            is_scan = text_len < 30 or (not raw_text and not has_tables)

            page_type = "text"
            extraction_method = "native"
            ocr_status = "none"
            confidence: float | None = None

            if is_scan:
                page_type = "scan"
                ocr_adapter = get_ocr_adapter()
                try:
                    if ocr_adapter.is_available():
                        # OCR 适配器需要完整 PDF 字节；仅在首次遇到扫描页时
                        # 惰性读取，保持普通文本 PDF 的流式解析特性。
                        if source_bytes is None:
                            current_pos = source.tell() if hasattr(source, "tell") else None
                            try:
                                if hasattr(source, "seek"):
                                    source.seek(0)
                                source_bytes = source.read()
                            finally:
                                if current_pos is not None and hasattr(source, "seek"):
                                    source.seek(current_pos)
                        ocr_result = ocr_adapter.ocr_pdf_page(source_bytes, idx - 1)
                        ocr_text = ocr_result.get("text", "")
                        confidence = ocr_result.get("confidence")
                        if ocr_text.strip():
                            cleaned = _clean_text(ocr_text)
                            raw_text = ocr_text
                            extraction_method = "ocr"
                            ocr_status = "success"
                            ocr_pages += 1
                            page_type = "scan" if text_len < 5 else "mixed"
                        else:
                            ocr_status = "failed"
                            failed_pages += 1
                            page_type = "scan"
                    else:
                        ocr_status = "failed"
                        failed_pages += 1
                        page_type = "scan"
                except Exception as exc:
                    logger.warning("ocr_page_failed", page=idx, error=str(exc))
                    ocr_status = "failed"
                    failed_pages += 1
                    page_type = "scan"
            elif has_tables:
                page_type = "mixed"

            # Markdown 文本（含表格）
            md_parts: list[str] = []
            if cleaned:
                md_parts.append(cleaned)
            for t in tables:
                md_parts.append(_table_to_markdown(t))
            markdown_text = "\n\n".join(md_parts)

            # 表格并入 chunk
            for t in tables:
                table_count += 1
                chunks.append(
                    ParsedChunk(
                        page_start=idx,
                        page_end=idx,
                        heading_path=" > ".join(heading_stack) if heading_stack else None,
                        content=_table_to_markdown(t),
                        chunk_type="table",
                        metadata={"tables": t},
                    )
                )

            pages.append(
                ParsedPage(
                    page_number=idx,
                    location_label=f"P{idx}",
                    raw_text=raw_text,
                    cleaned_text=cleaned,
                    markdown_text=markdown_text,
                    page_type=page_type,
                    extraction_method=extraction_method,
                    ocr_status=ocr_status,
                    confidence=confidence,
                    tables=tables,
                )
            )

            # 简单目录识别：短行 + 数字结尾 → 视为标题
            for line in cleaned.splitlines():
                s = line.strip()
                if 2 <= len(s) <= 60 and re.match(r"^[一-鿿A-Za-z0-9 一二三四五六七八九十.．、·（）()]+$", s):
                    # 检测标题层级
                    level = _detect_heading_level(s)
                    if level is not None and not re.search(r"^\d+[.．、)]?\s*$", s):
                        heading_text = re.sub(r"^\d+[.．、]?\s*", "", s).strip(" ")
                        toc.append(
                            {
                                "heading_path": " > ".join([*heading_stack, heading_text])
                                if heading_stack
                                else heading_text,
                                "heading_text": heading_text,
                                "level": level,
                                "page": idx,
                            }
                        )

    # 段落分块
    chunk_blocks = _split_paragraph_chunks(pages)
    chunks.extend(chunk_blocks)

    _assign_chunk_sequences(chunks)
    return ParsedDocument(
        pages=pages,
        chunks=chunks,
        toc=toc,
        total_pages=total,
        ocr_pages=ocr_pages,
        failed_pages=failed_pages,
        table_count=table_count,
    )


def _detect_heading_level(line: str) -> int | None:
    """简单标题层级识别。"""
    if re.match(r"^第[一二三四五六七八九十]+[章节部分]", line):
        return 1
    if re.match(r"^\d+\.\d+(\.\d+)?\s", line):
        return 2
    if re.match(r"^\d+[.．、)]\s", line):
        return 1
    # 短句且不含句号结尾
    if len(line) <= 30 and not line.endswith(("。", "；", "，")):
        return 3
    return None


def _split_paragraph_chunks(pages: list[ParsedPage]) -> list[ParsedChunk]:
    """将页面文本按段落切分为 chunk。"""
    chunks: list[ParsedChunk] = []
    for page in pages:
        if not page.cleaned_text:
            continue
        for para in re.split(r"\n\s*\n", page.cleaned_text):
            para = para.strip()
            if len(para) < 8:
                continue
            chunks.append(
                ParsedChunk(
                    page_start=page.page_number,
                    page_end=page.page_number,
                    heading_path=None,
                    content=para,
                    chunk_type="paragraph",
                )
            )
    return chunks


def _assign_chunk_sequences(chunks: list[ParsedChunk]) -> None:
    """为解析结果补齐稳定的文档顺序。"""
    for index, chunk in enumerate(chunks, start=1):
        chunk.sequence = index


def _table_to_json(table: list[list]) -> dict:
    """PDF 表格 → {headers: [], rows: []}。"""
    rows = []
    for r in table:
        rows.append([(cell or "").strip() for cell in r])
    if not rows:
        return {"headers": [], "rows": []}
    return {"headers": rows[0], "rows": rows[1:], "full_rows": rows}


def _table_to_markdown(table: dict | list) -> str:
    if isinstance(table, list):
        table = _table_to_json(table)
    headers = table.get("headers", [])
    rows = table.get("rows", [])
    lines = ["| " + " | ".join(headers) + " |", "|" + "---|" * len(headers)]
    for row in rows[:50]:
        lines.append("| " + " | ".join(str(c) for c in row) + " |")
    return "\n".join(lines)


# ============================================================
# DOCX 解析
# ============================================================

def _parse_docx(data: bytes) -> ParsedDocument:
    import io

    return _parse_docx_from_source(io.BytesIO(data))


def _parse_docx_from_source(source) -> ParsedDocument:
    """从文件路径或二进制流解析 DOCX。"""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    doc = Document(source)

    pages: list[ParsedPage] = []
    chunks: list[ParsedChunk] = []
    toc: list[dict] = []
    table_count = 0

    # DOCX 通常没有稳定的分页信息。优先使用 Word 写入的分页标记，
    # 没有标记时继续使用段落位置，避免伪装成真实页码。
    PAGE_SIZE = 40
    para_blocks: list[tuple[str, int | None, int, bool, dict[str, Any]]] = []
    page_no = 1
    sequence = 0
    has_explicit_page_break = False

    # 必须遍历 document.body 的子节点；分别读取 doc.paragraphs 和
    # doc.tables 会把表格全部移动到文档末尾，破坏施工方案原始顺序。
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                if _docx_block_has_page_break(child):
                    has_explicit_page_break = True
                    page_no += 1
                continue
            level = _detect_docx_heading_level(para)
            sequence += 1
            para_blocks.append((text, level, page_no, False, {"sequence": sequence}))
            if _docx_block_has_page_break(child):
                has_explicit_page_break = True
                page_no += 1
        elif isinstance(child, CT_Tbl):
            try:
                table = Table(child, doc)
                table_json = _table_to_json([[cell.text for cell in row.cells] for row in table.rows])
                table_md = _table_to_markdown(table_json)
                if table_md.strip():
                    table_count += 1
                    sequence += 1
                    para_blocks.append(
                        (table_md, -1, page_no, True, {"sequence": sequence, "tables": table_json})
                    )
                if _docx_block_has_page_break(child):
                    has_explicit_page_break = True
                    page_no += 1
            except Exception as exc:
                logger.warning("docx_table_parse_failed", error=str(exc))

    # 没有显式分页的DOCX仍使用稳定的虚拟页标签，仅用于定位，不称为真实PDF页。
    if page_no == 1 and para_blocks:
        para_blocks = [
            (text, level, (index - 1) // PAGE_SIZE + 1, is_table, metadata)
            for index, (text, level, _page, is_table, metadata) in enumerate(para_blocks, start=1)
        ]

    # 当前页文本
    page_texts: dict[int, list[str]] = {}
    for text, _level, page_no, _is_table, _metadata in para_blocks:
        page_texts.setdefault(page_no, []).append(text)

    # 构建页面
    heading_stack: list[str] = []
    for text, level, page_no, is_table, metadata in para_blocks:
        content = text

        if is_table:
            chunks.append(
                ParsedChunk(
                        page_start=page_no,
                        page_end=page_no,
                        heading_path=" > ".join(heading_stack) if heading_stack else None,
                        content=content,
                        chunk_type="table",
                        metadata=metadata,
                    )
            )
            continue

        if level is not None:
            # 维护标题栈
            clean_title = re.sub(r"^\d+[.．、]?\s*", "", text)
            if len(heading_stack) >= level:
                heading_stack = heading_stack[: level - 1]
            heading_stack.append(clean_title)
            toc.append(
                {
                    "heading_path": " > ".join(heading_stack),
                    "heading_text": clean_title,
                    "level": level,
                    "page": page_no,
                    "page_start": page_no,
                }
            )
            chunks.append(
                ParsedChunk(
                    page_start=page_no,
                    page_end=page_no,
                    heading_path=" > ".join(heading_stack),
                    content=text,
                    chunk_type="heading",
                    metadata=metadata,
                )
            )
        else:
            chunks.append(
                ParsedChunk(
                    page_start=page_no,
                    page_end=page_no,
                    heading_path=" > ".join(heading_stack) if heading_stack else None,
                    content=text,
                    chunk_type="paragraph",
                    metadata=metadata,
                )
            )

    for page_no, texts in sorted(page_texts.items()):
        raw = "\n".join(texts)
        cleaned = _clean_text(raw)
        pages.append(
            ParsedPage(
                page_number=page_no,
                location_label=f"段落{page_no}",
                raw_text=raw,
                cleaned_text=cleaned,
                markdown_text=cleaned,
                page_type="text",
                extraction_method="native",
                ocr_status="none",
                metadata={"is_virtual_page": not has_explicit_page_break},
            )
        )

    _assign_chunk_sequences(chunks)
    for chunk in chunks:
        chunk.metadata["is_virtual_page"] = not has_explicit_page_break
        chunk.metadata["location_label"] = (
            f"段落{chunk.page_start}" if not has_explicit_page_break else f"P{chunk.page_start}"
        )
    return ParsedDocument(
        pages=pages,
        chunks=chunks,
        toc=toc,
        total_pages=len(pages),
        ocr_pages=0,
        failed_pages=0,
        table_count=table_count,
    )


def _docx_block_has_page_break(block) -> bool:
    """检查Word正文块中的显式分页标记。"""
    xml = block.xml if hasattr(block, "xml") else str(block)
    return bool(
        re.search(r"lastRenderedPageBreak", xml)
        or re.search(r"w:br[^>]+w:type=[\"']page[\"']", xml)
    )


def _detect_docx_heading_level(para) -> int | None:
    """识别 DOCX 标题级别（样式名 或 数字编号）。"""
    style_name = (para.style.name or "").lower() if para.style else ""
    for i in range(1, 7):
        if f"heading {i}" in style_name or f"标题 {i}" in style_name:
            return i
    text = para.text.strip()
    if re.match(r"^第[一二三四五六七八九十]+[章节部分]", text):
        return 1
    if re.match(r"^\d+\.\d+\s", text):
        return 2
    if re.match(r"^\d+[.．、)]\s", text):
        return 1
    return None


# ============================================================
# TXT 解析
# ============================================================

def _parse_txt(data: bytes) -> ParsedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("gbk", errors="replace")

    pages: list[ParsedPage] = []
    chunks: list[ParsedChunk] = []
    toc: list[dict] = []
    heading_stack: list[str] = []

    lines = text.splitlines()
    PAGE_SIZE = 50  # 每 50 行虚拟一页
    total_pages = max(1, (len(lines) + PAGE_SIZE - 1) // PAGE_SIZE)

    current_para: list[str] = []
    current_page = 1
    para_start_page = 1

    def flush_para():
        nonlocal current_para, para_start_page
        if current_para:
            content = "\n".join(current_para)
            chunks.append(
                ParsedChunk(
                    page_start=para_start_page,
                    page_end=current_page,
                    heading_path=" > ".join(heading_stack) if heading_stack else None,
                    content=content,
                    chunk_type="paragraph",
                )
            )
            current_para = []

    for idx, line in enumerate(lines, start=1):
        page_no = (idx - 1) // PAGE_SIZE + 1
        s = line.strip()
        if not s:
            flush_para()
            continue

        level = _detect_heading_level(s)
        if level is not None and len(s) <= 60:
            flush_para()
            clean_title = re.sub(r"^\d+[.．、]?\s*", "", s).strip(" ")
            if len(heading_stack) >= level:
                heading_stack = heading_stack[: level - 1]
            heading_stack.append(clean_title)
            toc.append(
                {
                    "heading_path": " > ".join(heading_stack),
                    "heading_text": clean_title,
                    "level": level,
                    "page": page_no,
                    "page_start": page_no,
                }
            )
            chunks.append(
                ParsedChunk(
                    page_start=page_no,
                    page_end=page_no,
                    heading_path=" > ".join(heading_stack),
                    content=s,
                    chunk_type="heading",
                )
            )
            if not current_para:
                para_start_page = page_no
            continue

        if not current_para:
            para_start_page = page_no
        current_para.append(s)

    flush_para()

    # 构建页面
    page_lines: dict[int, list[str]] = {}
    for idx, line in enumerate(lines, start=1):
        page_no = (idx - 1) // PAGE_SIZE + 1
        page_lines.setdefault(page_no, []).append(line)

    for page_no in range(1, total_pages + 1):
        raw = "\n".join(page_lines.get(page_no, []))
        cleaned = _clean_text(raw)
        pages.append(
            ParsedPage(
                page_number=page_no,
                location_label=f"段落{page_no}",
                raw_text=raw,
                cleaned_text=cleaned,
                markdown_text=cleaned,
                page_type="text",
                extraction_method="native",
                ocr_status="none",
                metadata={"is_virtual_page": True, "source": "txt"},
            )
        )

    _assign_chunk_sequences(chunks)
    return ParsedDocument(
        pages=pages,
        chunks=chunks,
        toc=toc,
        total_pages=total_pages,
        ocr_pages=0,
        failed_pages=0,
        table_count=0,
    )


# ============================================================
# 统一入口
# ============================================================

def _parse_unknown_bytes(data: bytes) -> ParsedDocument:
    # 其他类型：仅存文本
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception:
        text = ""
    return ParsedDocument(
        pages=[
            ParsedPage(
                page_number=1,
                location_label="P1",
                raw_text=text,
                cleaned_text=_clean_text(text),
                markdown_text=_clean_text(text),
                page_type="text",
                extraction_method="native",
                ocr_status="none",
            )
        ],
        chunks=[],
        toc=[],
        total_pages=1,
        ocr_pages=0,
        failed_pages=0,
        table_count=0,
    )


def parse_document_bytes(data: bytes, file_type: str) -> ParsedDocument:
    """按文件类型解析（小文件/测试使用，整文件载入内存）。"""
    if file_type == "pdf":
        return _parse_pdf(data)
    if file_type == "docx":
        return _parse_docx(data)
    if file_type == "txt":
        return _parse_txt(data)
    return _parse_unknown_bytes(data)


def parse_document_path(path: str | Path, file_type: str) -> ParsedDocument:
    """从磁盘文件路径解析，供大文件使用：PDF/DOCX 以文件流方式打开，避免整文件读入内存。"""
    path = str(path)
    if file_type == "pdf":
        with open(path, "rb") as fh:
            return _parse_pdf_from_source(fh)
    if file_type == "docx":
        with open(path, "rb") as fh:
            return _parse_docx_from_source(fh)
    if file_type == "txt":
        # 文本文件通常较小；逐块尝试解码
        with open(path, "rb") as fh:
            head = fh.read(4096)
        try:
            head.decode("utf-8")
            encoding = "utf-8"
        except UnicodeDecodeError:
            encoding = "gbk"
        with open(path, encoding=encoding, errors="replace") as fh:
            return _parse_txt(fh.read().encode("utf-8"))
    with open(path, "rb") as fh:
        return _parse_unknown_bytes(fh.read())


def compute_sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def get_all_doc_types() -> dict[str, str]:
    return DOC_TYPES
