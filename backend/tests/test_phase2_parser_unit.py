"""解析器单元测试：PDF/DOCX 段落表格、扫描页 OCR 降级。"""

from __future__ import annotations

import io
import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

os.environ.setdefault("DATABASE_URL", "sqlite:////tmp/fastvideo_parser_test.db")
os.environ.setdefault("USE_CELERY", "false")

from app.services.document_parser import parse_document_bytes  # noqa: E402
from app.adapters.ocr import MockOCRAdapter  # noqa: E402


def _make_pdf_with_text() -> bytes:
    """生成一个含文本和表格的 PDF。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("一、项目概况", styles["Heading1"]),
        Paragraph("项目名称：单元测试项目", styles["Normal"]),
        Paragraph("总建筑面积 52800 平方米，总工期 540 日历天。", styles["Normal"]),
        Table(
            [["参数", "数值", "单位"], ["建筑面积", "52800", "㎡"], ["总工期", "540", "日历天"]],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]),
        ),
    ]
    doc.build(story)
    return buf.getvalue()


def _make_docx() -> bytes:
    """生成一个含标题/段落/表格的 DOCX。"""
    from docx import Document

    doc = Document()
    doc.add_heading("施工组织设计", level=1)
    doc.add_paragraph("本工程采用流水施工组织。")
    doc.add_heading("施工部署", level=2)
    doc.add_paragraph("成立项目管理部，配置专职管理人员。")
    table = doc.add_table(rows=3, cols=3)
    cells = [
        ("阶段", "内容", "工期"),
        ("基础", "桩基及基坑", "120天"),
        ("主体", "结构封顶", "360天"),
    ]
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            table.cell(i, j).text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_pdf_parse_pages_and_tables():
    data = _make_pdf_with_text()
    parsed = parse_document_bytes(data, "pdf")
    assert parsed.total_pages >= 1
    assert len(parsed.pages) >= 1
    # 应提取出参数
    all_text = "\n".join(p.cleaned_text or "" for p in parsed.pages)
    assert "52800" in all_text or "540" in all_text
    # 表格
    assert parsed.table_count >= 1


def test_pdf_scan_page_ocr_degradation():
    """纯图片页（无文本）应被标记为 scan，OCR 失败不抛异常。"""
    # 构造一张空白页（无文本）
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFillColorRGB(1, 1, 1)
    c.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
    c.showPage()
    c.save()
    data = buf.getvalue()

    parsed = parse_document_bytes(data, "pdf")
    assert parsed.total_pages == 1
    assert parsed.pages[0].page_type == "scan"
    # OCR 状态应为 success（Mock/真实OCR识别到文本）或 failed（OCR不可用/无文本）
    # 关键：OCR 失败不得导致解析抛异常，页面状态应保留
    assert parsed.pages[0].ocr_status in ("success", "failed")
    # 文档级统计
    assert parsed.failed_pages >= 0
    assert parsed.ocr_pages >= 0


def test_mock_ocr_adapter():
    adapter = MockOCRAdapter()
    assert adapter.is_available() is True
    result = adapter.ocr_pdf_page(b"dummy", 0)
    assert "text" in result
    assert "confidence" in result


def test_docx_parse_headings_and_tables():
    data = _make_docx()
    parsed = parse_document_bytes(data, "docx")
    assert parsed.total_pages >= 1
    # 目录识别
    assert len(parsed.toc) >= 2, f"应识别出标题，实际{parsed.toc}"
    # 表格
    assert parsed.table_count >= 1
    # 虚拟分页标记
    assert parsed.pages[0].metadata.get("is_virtual_page") is True
    # 标题层级路径
    heading_chunks = [c for c in parsed.chunks if c.chunk_type == "heading"]
    assert len(heading_chunks) >= 2


def test_docx_keeps_paragraph_table_paragraph_order_and_sequences():
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一段：先完成场地清理。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "工序"
    table.cell(0, 1).text = "基础"
    table.cell(1, 0).text = "参数"
    table.cell(1, 1).text = "分层开挖"
    doc.add_paragraph("第三段：随后进入基础施工。")
    buf = io.BytesIO()
    doc.save(buf)

    parsed = parse_document_bytes(buf.getvalue(), "docx")
    chunks = [chunk for chunk in parsed.chunks if chunk.content]
    sequences = [chunk.sequence for chunk in chunks]
    assert sequences == sorted(sequences)
    assert len(sequences) == len(set(sequences))
    table_index = next(i for i, chunk in enumerate(chunks) if chunk.chunk_type == "table")
    first_index = next(i for i, chunk in enumerate(chunks) if "第一段" in chunk.content)
    last_index = next(i for i, chunk in enumerate(chunks) if "第三段" in chunk.content)
    assert first_index < table_index < last_index


def test_txt_parse():
    content = "一、项目概况\n项目名称：测试项目\n建筑面积 10000 平方米。\n二、施工部署\n内容。\n"
    parsed = parse_document_bytes(content.encode("utf-8"), "txt")
    assert parsed.total_pages >= 1
    assert len(parsed.toc) >= 2, "TXT 应识别标题"
    all_text = "\n".join(p.cleaned_text or "" for p in parsed.pages)
    assert "10000" in all_text
